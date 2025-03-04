import re
import os
import docx

def fix_docx_loader():
    """Create a fixed DOCX loader for numbered options"""
    print("Creating a fixed DOCX loader...")
    
    # Load the numbers_test.docx file which contains numbered options
    # First let's create a test file with numbered options
    doc = docx.Document()
    
    # Question 1 with numbered options
    doc.add_paragraph('What is the capital of Italy?')
    
    # Option 1 - not bold
    p = doc.add_paragraph()
    p.add_run('1. Madrid')
    
    # Option 2 - not bold
    p = doc.add_paragraph()
    p.add_run('2. Berlin')
    
    # Option 3 - bold (correct answer)
    p = doc.add_paragraph()
    p.add_run('3. ')
    p.add_run('Rome').bold = True
    
    # Option 4 - not bold
    p = doc.add_paragraph()
    p.add_run('4. Athens')
    
    # Feedback
    doc.add_paragraph('Answer Feedback: Rome is the capital of Italy since 1871.')
    doc.add_paragraph()  # Blank line
    
    # Save the document
    os.makedirs("quizzes", exist_ok=True)
    doc.save('quizzes/numbers_test.docx')
    print("Created quizzes/numbers_test.docx with numbered options")
    
    # Now let's create a version of the _load_from_docx function that handles numbered options
    with open('numbered_option_handler.py', 'w') as f:
        f.write('''
def _load_from_docx(self, filename):
    """Load questions from a DOCX file, preserving formatting for both letter and numbered options."""
    doc = docx.Document(filename)
    
    # We'll process paragraphs sequentially
    current_question = None
    current_options = []
    current_feedback = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Empty paragraph means end of current question
            if current_question:
                question_type = "free-response"
                if current_options:
                    question_type = "multiple-choice"
                elif current_question.lower().find("true/false") > -1:
                    question_type = "true-false"
                    current_options = ["True", "False"]
                
                self.add_question(Question(current_question, question_type, current_options, current_feedback))
                current_question = None
                current_options = []
                current_feedback = None
        elif text.startswith("Answer Feedback:"):
            # This is feedback for the current question
            current_feedback = text
        # Check for letter-style options (a), b), etc.)
        elif any(text.startswith(f"{chr(97 + i)})") for i in range(26)):
            # This is a letter option - check for bold parts
            option_text = text
            
            # Explicitly check each run for bold property
            has_bold = False
            
            for run in para.runs:
                if run.bold is True:  # Must explicitly check against True
                    has_bold = True
                    bold_text = run.text
                    # Mark the bold part with markdown
                    option_text = option_text.replace(bold_text, f"**{bold_text}**")
                    break
            
            current_options.append(option_text)
        # Check for number-style options (1., 2., etc.)
        elif re.match(r'^\\d+\\.\\s+.+', text):
            # This is a numbered option
            number_match = re.match(r'^(\\d+)\\.\\s+(.+)$', text)
            if number_match:
                number = int(number_match.group(1))
                content = number_match.group(2)
                
                # Convert to letter format (1 -> a, 2 -> b, etc.)
                letter = chr(96 + number) if 1 <= number <= 26 else chr(96 + (number % 26))
                letter_prefix = f"{letter}) "
                
                # Check for bold formatting
                has_bold = False
                for run in para.runs:
                    if run.bold is True:
                        has_bold = True
                        # Find which part of the text is the bold part
                        bold_text = run.text
                        if bold_text in content:
                            # Mark content as bold
                            modified_content = content.replace(bold_text, f"**{bold_text}**")
                            option_text = f"{letter_prefix}{modified_content}"
                            break
                
                # If no bold parts found, use plain formatting
                if not has_bold:
                    option_text = f"{letter_prefix}{content}"
                
                current_options.append(option_text)
        elif not current_question:
            # This must be a new question text
            current_question = text
    
    # Add the last question if there is one
    if current_question:
        question_type = "free-response"
        if current_options:
            question_type = "multiple-choice"
        elif current_question.lower().find("true/false") > -1:
            question_type = "true-false"
            current_options = ["True", "False"]
        
        self.add_question(Question(current_question, question_type, current_options, current_feedback))
''')
    
    print("Created numbered_option_handler.py with updated loader function")
    print("To use this function, copy and paste it into the QuestionBank class in quiz_randomizer.py")

if __name__ == "__main__":
    fix_docx_loader()
