import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import sys
import random
import re
from fpdf import FPDF
from typing import List, Dict, Set, Tuple
import unicodedata
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx2txt
from striprtf.striprtf import rtf_to_text
from bold_formatter import create_quiz_from_docx
from config_manager import ConfigManager

# Helper function to locate resources when packaged with PyInstaller
def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(os.path.dirname(__file__))
    return os.path.join(base_path, relative_path)

# Custom FPDF class with UTF-8 support
class UTF8PDF(FPDF):
    def __init__(self):
        super().__init__()
        # Set up the default font that supports extended characters
        self.add_font('DejaVu', '', resource_path('DejaVuSansCondensed.ttf'), uni=True)
        self.add_font('DejaVu', 'B', resource_path('DejaVuSansCondensed-Bold.ttf'), uni=True)
    
    def sanitize_text(self, text):
        """Sanitize text to avoid encoding issues"""
        # Replace any problematic characters or convert to ASCII equivalent
        if isinstance(text, str):
            # ASCII-only approximation (removes accents but keeps base letters)
            text = unicodedata.normalize('NFKD', text)
            text = ''.join([c for c in text if not unicodedata.combining(c)])
            # Remove any remaining non-ASCII characters
            text = text.encode('ascii', 'replace').decode('ascii')
        return text


class Question:
    """Class to represent a quiz question with its text and options."""
    
    def __init__(self, text: str, question_type: str = "unknown", options: List[str] = None, feedback: str = None):
        self.text = text.strip()
        self.question_type = question_type  # "multiple-choice", "true-false", or "free-response"
        self.options = options or []
        self.feedback = feedback  # Store answer feedback
        self.id = hash(self.text)  # Simple way to identify unique questions
        
        # Track which option is correct (for preserving order)
        self.correct_option_index = -1
        for i, option in enumerate(self.options):
            if "**" in option:
                self.correct_option_index = i
                break
    
    def __str__(self) -> str:
        result = self.text + "\n"  # Add newline after question text
        if self.options:
            for i, option in enumerate(self.options):
                # Check if the option already has a prefix (a), b), etc.)
                if not re.match(r'^[a-z]\)\s+', option.lower()):
                    # If no prefix, add one
                    result += f"{chr(97 + i)}) {option}\n"
                else:
                    # If it has a prefix, use it as is
                    result += f"{option}\n"
        if self.feedback:
            result += f"{self.feedback}\n"
        return result.rstrip()  # Remove trailing newline


class QuestionBank:
    """Class to manage a collection of questions."""
    
    def __init__(self):
        self.questions: List[Question] = []
    
    def add_question(self, question: Question) -> None:
        """Add a question to the bank."""
        self.questions.append(question)
    
    def get_size(self) -> int:
        """Get the number of questions in the bank."""
        return len(self.questions)
    
    def load_from_file(self, filename: str) -> None:
        """Load questions from a file (supports .txt, .docx, and .rtf)."""
        if not os.path.exists(filename):
            raise FileNotFoundError(f"Question bank file {filename} not found")
        
        # Determine file type and extract content accordingly
        file_extension = os.path.splitext(filename)[1].lower()
        
        if file_extension == '.docx':
            # For DOCX files, we'll use a specialized approach to preserve formatting
            self._load_from_docx(filename)
            return
        
        elif file_extension == '.rtf':
            # Extract text from RTF file
            with open(filename, 'r', encoding='utf-8', errors='ignore') as file:
                rtf_content = file.read()
            content = rtf_to_text(rtf_content)
            # Note: striprtf doesn't preserve formatting like bold
            # We would need a more sophisticated RTF parser for that
        
        else:  # Default to txt
            with open(filename, 'r', encoding='utf-8') as file:
                content = file.read()
        
        # Process text-based content (TXT or RTF)
        # Split content into individual questions
        question_blocks = re.split(r'\n\s*\n', content)
        
        for block in question_blocks:
            if not block.strip():
                continue
            
            lines = block.strip().split('\n')
            if not lines:  # Skip empty blocks
                continue
                
            question_text = lines[0]
            
            # Try to determine question type
            question_type = "free-response"  # Default
            options = []
            feedback = None
            
            # Check for answer feedback
            feedback_pattern = re.compile(r'^Answer\s+Feedback:\s+.+')
            feedback_lines = [line for line in lines if feedback_pattern.match(line)]
            if feedback_lines:
                feedback = feedback_lines[0]
                # Remove feedback line from lines to process
                lines = [line for line in lines if not feedback_pattern.match(line)]
            
            # Check for multiple choice format - support both letter (a), b)) and numbered formats (1., 2.)
            letter_option_pattern = re.compile(r'^[a-z]\)\s+.+')
            number_option_pattern = re.compile(r'^\d+\.\s+.+')
            
            has_letter_options = any(letter_option_pattern.match(line) for line in lines[1:])
            has_number_options = any(number_option_pattern.match(line) for line in lines[1:])
            has_options = has_letter_options or has_number_options
            
            if has_options:
                question_type = "multiple-choice"
                options = []
                for line in lines[1:]:
                    # Check for letter-style options (a), b), etc.)
                    if letter_option_pattern.match(line):
                        options.append(line)
                    # Check for number-style options (1., 2., etc.)
                    elif number_option_pattern.match(line):
                        # Convert numbered format to letter format for consistency
                        number_match = re.match(r'^(\d+)\.\s+(.+)$', line)
                        if number_match:
                            number = int(number_match.group(1))
                            option_text = number_match.group(2)
                            # Convert to letter format (1 -> a, 2 -> b, etc.)
                            letter = chr(96 + number) if 1 <= number <= 26 else chr(96 + (number % 26))
                            converted_option = f"{letter}) {option_text}"
                            options.append(converted_option)
                
            # Check for true/false
            elif any(re.search(r'true\s*\/\s*false', line.lower()) for line in lines):
                question_type = "true-false"
                options = ["True", "False"]
            
            self.add_question(Question(question_text, question_type, options, feedback))
    
    def _load_from_docx(self, filename: str) -> None:
        """Load questions from a DOCX file, preserving formatting."""
        print("\nDEBUG: Starting to load from DOCX")
        doc = docx.Document(filename)
        
        # We'll process paragraphs sequentially
        current_question = None
        current_options = []
        current_feedback = None
        processing_options = False  # Flag to track if we're processing options
        option_number = 0  # Counter for tracking which option we're processing
        expected_option_pattern = None  # Track what kind of options we're expecting (numbered or lettered)
        question_count = 0  # Track how many questions we've processed
        
        def save_current_question():
            """Helper function to save the current question and reset state"""
            nonlocal current_question, current_options, current_feedback, processing_options, option_number, expected_option_pattern, question_count
            if current_question and current_options:  # Only save if we have both question and options
                question_count += 1
                print(f"\nDEBUG: Saving question #{question_count}:")
                print(f"  Question text: {current_question}")
                print(f"  Number of options: {len(current_options)}")
                print(f"  Options:")
                for i, opt in enumerate(current_options, 1):
                    print(f"    {i}. {opt}")
                print(f"  Feedback: {current_feedback}")
                print("-" * 80)
                
                # Determine question type
                question_type = "free-response"  # Default type
                if current_question.lower().find("true/false") > -1:
                    question_type = "true-false"
                    current_options = ["True", "False"]
                elif current_options:
                    question_type = "multiple-choice"
                
                self.add_question(Question(current_question, question_type, current_options, current_feedback))
                # Reset state
                current_question = None
                current_options = []
                current_feedback = None
                processing_options = False
                option_number = 0
                expected_option_pattern = None
        
        def looks_like_question(text: str) -> bool:
            """Helper function to determine if text looks like a question"""
            # Don't consider very short text as questions
            if len(text) < 20:
                return False
            
            # Don't consider text that starts with typical option patterns
            if re.match(r'^[a-z0-9][\.\)]\s+', text.lower()):
                return False
            
            # Consider text that ends with a question mark as a question
            if text.endswith('?'):
                return True
            
            # Consider text that starts with question-like phrases
            question_starters = ['what', 'which', 'who', 'where', 'when', 'why', 'how', 'describe', 'explain', 'discuss']
            if any(text.lower().startswith(starter) for starter in question_starters):
                return True
            
            # Consider text that contains question-like phrases
            if any(f" {starter} " in text.lower() for starter in question_starters):
                return True
            
            return False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            print(f"\nDEBUG: Processing paragraph: '{text}'")
            if text:
                print(f"DEBUG: Paragraph runs:")
                for i, run in enumerate(para.runs):
                    print(f"  Run {i}: text='{run.text}', bold={run.bold}")
            
            # Skip empty paragraphs
            if not text:
                # If we were processing options and have both question and some options,
                # this empty line might indicate the end of the question
                if processing_options and current_question and len(current_options) > 0:
                    save_current_question()
                continue
            
            # Handle feedback
            if text.startswith("Answer Feedback:"):
                current_feedback = text
                if current_question and current_options:  # If we have a complete question
                    save_current_question()
                continue
            
            # Check if this looks like a new question
            is_question = looks_like_question(text)
            
            # Check if this is an option
            is_letter_option = bool(re.match(r'^[a-z]\)\s+.+', text.lower()))
            is_number_option = bool(re.match(r'^\d+\.\s+.+', text))
            is_option = is_letter_option or is_number_option
            
            print(f"DEBUG: Analysis - is_question: {is_question}, is_option: {is_option}")
            
            if is_question and not is_option:
                # If we were already processing a question with options, save it
                if current_question and current_options:
                    save_current_question()
                
                # Start new question
                current_question = text
                processing_options = True
                option_number = 0
                expected_option_pattern = None
                print(f"DEBUG: Found new question: {current_question}")
                
                # Check if it's a true/false question
                if "true/false" in current_question.lower():
                    print("DEBUG: Detected True/False question")
                    current_options = ["True", "False"]
                    save_current_question()
                    continue
            
            # Handle letter-style options (a), b), etc.)
            elif is_letter_option:
                print("DEBUG: Processing letter-style option")
                if expected_option_pattern is None:
                    expected_option_pattern = "letter"
                elif expected_option_pattern != "letter":
                    # If we were expecting numbered options but got a letter option,
                    # this might be a new question that our detection missed
                    if current_question and current_options:
                        save_current_question()
                    current_question = text
                    processing_options = True
                    option_number = 0
                    expected_option_pattern = None
                    continue
                
                processing_options = True
                option_number += 1
                
                option_text = text.split(')', 1)
                letter_part = option_text[0] + ')'
                content_part = option_text[1].strip() if len(option_text) > 1 else ''
                
                formatted_content = self._process_bold_text(para.runs)
                if not formatted_content:
                    formatted_content = content_part
                
                option = f"{letter_part} {formatted_content}"
                print(f"DEBUG: Adding letter option ({option_number}): {option}")
                current_options.append(option)
            
            # Handle numbered options (1., 2., etc.)
            elif is_number_option:
                print("DEBUG: Processing numbered option")
                if expected_option_pattern is None:
                    expected_option_pattern = "number"
                elif expected_option_pattern != "number":
                    # If we were expecting letter options but got a numbered option,
                    # this might be a new question that our detection missed
                    if current_question and current_options:
                        save_current_question()
                    current_question = text
                    processing_options = True
                    option_number = 0
                    expected_option_pattern = None
                    continue
                
                processing_options = True
                option_number += 1
                
                # Process bold text first
                formatted_content = self._process_bold_text(para.runs)
                if not formatted_content:
                    formatted_content = text.split('.', 1)[1].strip()
                
                # Convert to letter format
                letter = chr(96 + option_number)
                option = f"{letter}) {formatted_content}"
                print(f"DEBUG: Adding numbered option ({option_number}): {option}")
                current_options.append(option)
            
            # If we're processing options and this looks like an unlabeled option
            elif processing_options and not current_feedback and not is_question:
                print("DEBUG: Processing unlabeled option")
                option_number += 1
                
                formatted_content = self._process_bold_text(para.runs)
                if not formatted_content:
                    formatted_content = text
                
                letter = chr(96 + option_number)
                option = f"{letter}) {formatted_content}"
                print(f"DEBUG: Adding unlabeled option ({option_number}): {option}")
                current_options.append(option)
            else:
                print(f"DEBUG: Unhandled paragraph: {text}")
                # If this looks like it could be a question but we missed it
                if len(text) > 20 and not text.startswith("Answer Feedback:"):
                    print("DEBUG: Reconsidering as potential question...")
                    if current_question and current_options:
                        save_current_question()
                    current_question = text
                    processing_options = True
                    option_number = 0
                    expected_option_pattern = None
        
        # Save the last question if there is one
        if current_question and current_options:
            save_current_question()
        
        print(f"\nDEBUG: Finished loading questions. Total questions loaded: {question_count}")
        if question_count != 30:
            print("\nWARNING: Expected 30 questions but found {question_count}. Some questions may have been missed.")
    
    def _process_bold_text(self, runs) -> str:
        """Helper method to process bold text in paragraph runs."""
        formatted_content = ''
        for run in runs:
            if run.bold is True:  # Must explicitly check against True
                run_text = run.text.strip()
                if run_text:
                    formatted_content += f"*{run_text}*"
            else:
                formatted_content += run.text
        return formatted_content.strip()


class Quiz:
    """Class to represent a quiz with questions and metadata."""
    
    def __init__(self, title: str, questions: List[Question] = None):
        self.title = title
        self.questions = questions or []
    
    def add_question(self, question: Question) -> None:
        """Add a question to the quiz."""
        self.questions.append(question)
    
    def get_question_count(self) -> int:
        """Get the number of questions in the quiz."""
        return len(self.questions)
    
    def __str__(self) -> str:
        """String representation of the quiz for text output."""
        result = f"{self.title}\n\n"
        for i, question in enumerate(self.questions, 1):
            result += f"{i}. {question}\n\n"
        return result
    
    def to_docx(self, filename: str) -> None:
        """Export the quiz to a DOCX file that preserves formatting."""
        doc = docx.Document()
        temp_filename = filename + '.temp'
        
        # Add title if it exists
        if self.title:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(self.title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()  # Add some space
        
        # Add each question with its options and feedback
        for question in self.questions:
            # Add the question text
            doc.add_paragraph(question.text)
            
            # Add the options
            for option in question.options:
                # If option includes single asterisk style bold (*text*), we need to handle it
                bold_pattern = re.compile(r'\*(.+?)\*')
                bold_match = bold_pattern.search(option)
                
                if bold_match:
                    # Split into letter part and content part
                    parts = option.split(')', 1)
                    letter_part = parts[0] + ')'
                    content = parts[1].strip() if len(parts) > 1 else ''
                    
                    # Create a paragraph for this option
                    p = doc.add_paragraph()
                    
                    # Add the letter part
                    p.add_run(letter_part + ' ')
                    
                    # Process the content part, handling bold sections
                    pos = 0
                    for match in bold_pattern.finditer(content):
                        # Add any text before the bold part
                        if match.start() > pos:
                            p.add_run(content[pos:match.start()])
                        
                        # Add the bold text
                        bold_run = p.add_run(match.group(1))
                        bold_run.bold = True
                        
                        pos = match.end()
                    
                    # Add any remaining text after the last bold part
                    if pos < len(content):
                        p.add_run(content[pos:])
                else:
                    # Regular option without bold formatting
                    doc.add_paragraph(option)
            
            # Add the feedback if it exists
            if question.feedback:
                feedback_para = doc.add_paragraph()
                feedback_run = feedback_para.add_run(question.feedback)
                feedback_run.italic = True
                feedback_run.font.size = Pt(10)
            
            # Add a blank line between questions
            doc.add_paragraph()
        
        doc.save(temp_filename)
        
        # Use our specialized function to properly format the output
        create_quiz_from_docx(
            input_file=temp_filename,
            output_file=filename,
            title=self.title
        )
        
        # Clean up the temporary file
        try:
            os.remove(temp_filename)
        except:
            pass
    
    def to_pdf(self, filename: str) -> None:
        """Export the quiz to a PDF file."""
        pdf = UTF8PDF()
        pdf.add_page()
        
        # Set up fonts
        pdf.set_font("DejaVu", "", 12)
        
        # Add title if it exists
        if self.title:
            pdf.set_font("DejaVu", "B", 16)
            pdf.cell(0, 10, self.title, 0, 1, 'C')
            pdf.ln(5)
            pdf.set_font("DejaVu", "", 12)
        
        # Add each question with its options
        for i, question in enumerate(self.questions, 1):
            # Add question number and text
            pdf.set_font("DejaVu", "", 12)
            pdf.set_x(10)
            pdf.multi_cell(0, 8, f"{i}. {question.text}")
            
            pdf.ln(5)
            for j, option in enumerate(question.options):
                pdf.set_x(20)  # Indent options
                
                # Check if this option has bold text (single asterisk format)
                bold_pattern = re.compile(r'\*(.+?)\*')
                bold_match = bold_pattern.search(option)
                
                # Process option text to handle bold formatting
                if bold_match:
                    # Split into letter part and content part
                    parts = option.split(')', 1)
                    letter_part = parts[0] + ')'
                    content = parts[1].strip() if len(parts) > 1 else ''
                    
                    # Add the letter part
                    pdf.set_font("DejaVu", "", 12)
                    pdf.write(8, letter_part + ' ')
                    
                    # Process the content part, handling bold sections
                    pos = 0
                    for match in bold_pattern.finditer(content):
                        # Add any text before the bold part
                        if match.start() > pos:
                            pdf.set_font("DejaVu", "", 12)
                            pdf.write(8, content[pos:match.start()])
                        
                        # Add the bold text
                        pdf.set_font("DejaVu", "B", 12)
                        pdf.write(8, match.group(1))
                        pos = match.end()
                    
                    # Add any remaining text after the last bold part
                    if pos < len(content):
                        pdf.set_font("DejaVu", "", 12)
                        pdf.write(8, content[pos:])
                    
                    pdf.ln()
                else:
                    # Regular option without bold formatting
                    pdf.set_font("DejaVu", "", 12)
                    pdf.multi_cell(0, 8, option, 0, 'L')
            
            # Add the feedback if it exists
            if question.feedback:
                pdf.ln(2)
                pdf.set_font("DejaVu", "", 10)
                pdf.set_text_color(100, 100, 100)  # Gray color for feedback
                pdf.set_x(20)
                pdf.multi_cell(0, 6, question.feedback)
                pdf.set_text_color(0, 0, 0)  # Reset to black
            
            pdf.ln(8)  # Space between questions
        
        # Save the PDF
        pdf.output(filename)


class QuizRandomizer:
    """Main class to generate randomized quizzes from a question bank."""
    
    def __init__(self, question_bank: QuestionBank):
        self.question_bank = question_bank
    
    def create_quizzes(
        self, 
        num_quizzes: int, 
        questions_per_quiz: int, 
        allow_duplicates: bool = False,
        output_format: str = "pdf",
        output_dir: str = "quizzes"
    ) -> Tuple[List[Quiz], Dict]:
        """
        Create multiple quizzes with randomized questions.
        
        Args:
            num_quizzes: Number of quizzes to create
            questions_per_quiz: Number of questions per quiz
            allow_duplicates: If True, same question can appear in multiple quizzes
            output_format: Format to output quizzes ("text", "pdf")
            output_dir: Directory to save output files
            
        Returns:
            Tuple containing list of Quiz objects and metadata dictionary
        """
        total_questions_needed = num_quizzes * questions_per_quiz
        bank_size = self.question_bank.get_size()
        
        # Validate that we have enough unique questions if duplicates aren't allowed
        if not allow_duplicates and total_questions_needed > bank_size:
            raise ValueError(
                f"Not enough unique questions. Requested {total_questions_needed} but only have {bank_size}. "
                f"Either reduce the number of quizzes/questions per quiz or allow duplicates."
            )
        
        # Create quizzes
        quizzes = []
        metadata = {
            "num_quizzes": num_quizzes,
            "questions_per_quiz": questions_per_quiz,
            "total_questions_used": 0,
            "quizzes_with_extra_questions": [],
            "duplicate_stats": {}
        }
        
        # Determine distribution of questions
        if allow_duplicates:
            # Simple case: we just randomly select with replacement for each quiz
            for i in range(num_quizzes):
                quiz = Quiz(f"Quiz {i+1}")
                selected_questions = random.choices(self.question_bank.questions, k=questions_per_quiz)
                
                for question in selected_questions:
                    # Create a deep copy to avoid modifying the original
                    quiz.add_question(question)
                
                quizzes.append(quiz)
                
            # Track duplicates for metadata
            question_usage = {}
            for quiz in quizzes:
                for question in quiz.questions:
                    if question.id not in question_usage:
                        question_usage[question.id] = 0
                    question_usage[question.id] += 1
            
            for question_id, count in question_usage.items():
                if count > 1:
                    metadata["duplicate_stats"][question_id] = count
            
            metadata["total_questions_used"] = len(question_usage)
                
        else:
            # No duplicates allowed - select a subset of questions for all quizzes
            all_questions = self.question_bank.questions.copy()
            random.shuffle(all_questions)
            
            # Select only the number of questions needed (num_quizzes * questions_per_quiz)
            # or all questions if there aren't enough
            selected_questions = all_questions[:min(total_questions_needed, bank_size)]
            metadata["total_questions_used"] = len(selected_questions)
            
            # Distribute questions across quizzes
            for i in range(num_quizzes):
                quiz = Quiz(f"Quiz {i+1}")
                
                # Calculate start and end index for this quiz's questions
                start_idx = i * questions_per_quiz
                end_idx = min(start_idx + questions_per_quiz, len(selected_questions))
                
                # If we reach the end of available questions, we won't have a full quiz
                if start_idx >= len(selected_questions):
                    break
                
                # Add the questions to this quiz
                for j in range(start_idx, end_idx):
                    quiz.add_question(selected_questions[j])
                
                quizzes.append(quiz)
        
        # Create output directory if it doesn't exist
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Export quizzes to the specified format
        for i, quiz in enumerate(quizzes):
            try:
                if output_format == "pdf":
                    quiz.to_pdf(os.path.join(output_dir, f"quiz_{i+1}.pdf"))
                elif output_format == "docx":
                    quiz.to_docx(os.path.join(output_dir, f"quiz_{i+1}.docx"))
                else:  # default to text
                    with open(os.path.join(output_dir, f"quiz_{i+1}.txt"), 'w', encoding='utf-8') as f:
                        f.write(str(quiz))
            except Exception as e:
                # If preferred format fails, fall back to text format
                with open(os.path.join(output_dir, f"quiz_{i+1}.txt"), 'w', encoding='utf-8') as f:
                    f.write(str(quiz))
                print(f"Warning: Exported quiz {i+1} as text due to issues: {str(e)}")
        
        return quizzes, metadata


class SimpleQuizRandomizerGUI:
    """A simplified GUI for the Quiz Randomizer with no threading."""
    
    def __init__(self, root):
        self.root = root
        self.config = ConfigManager()
        
        # Use config values
        self.root.title(self.config.get_setting("gui", "title"))
        geometry = f"{self.config.get_setting('gui', 'window_width')}x{self.config.get_setting('gui', 'window_height')}"
        self.root.geometry(geometry)
        
        # Initialize with default values from config
        self.num_quizzes_var = tk.IntVar(value=self.config.get_setting("quiz_defaults", "num_quizzes"))
        self.questions_per_quiz_var = tk.IntVar(value=self.config.get_setting("quiz_defaults", "questions_per_quiz"))
        self.format_var = tk.StringVar(value=self.config.get_setting("quiz_defaults", "output_format"))
        self.output_dir_var = tk.StringVar(value=self.config.get_setting("quiz_defaults", "output_directory"))
        
        self.question_bank = None
        self.create_widgets()
    
    def create_widgets(self):
        # Main frame
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # File selection
        file_frame = ttk.LabelFrame(main_frame, text="Question Bank")
        file_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.file_var = tk.StringVar()
        file_entry = ttk.Entry(file_frame, textvariable=self.file_var, width=50)
        file_entry.grid(row=0, column=0, padx=5, pady=5)
        
        browse_btn = ttk.Button(file_frame, text="Browse...", command=self.browse_file)
        browse_btn.grid(row=0, column=1, padx=5, pady=5)
        
        load_btn = ttk.Button(file_frame, text="Load Questions", command=self.load_questions)
        load_btn.grid(row=1, column=0, columnspan=2, pady=5)
        
        self.status_var = tk.StringVar(value="No question bank loaded")
        status_label = ttk.Label(file_frame, textvariable=self.status_var)
        status_label.grid(row=2, column=0, columnspan=2, pady=5, sticky="w")
        
        # Quiz options
        options_frame = ttk.LabelFrame(main_frame, text="Quiz Options")
        options_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # Number of quizzes
        ttk.Label(options_frame, text="Number of quizzes:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Spinbox(options_frame, from_=1, to=100, textvariable=self.num_quizzes_var, width=5).grid(row=0, column=1, padx=5, pady=5)
        
        # Questions per quiz
        ttk.Label(options_frame, text="Questions per quiz:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        ttk.Spinbox(options_frame, from_=1, to=100, textvariable=self.questions_per_quiz_var, width=5).grid(row=1, column=1, padx=5, pady=5)
        
        # Output format
        ttk.Label(options_frame, text="Output format:").grid(row=0, column=2, sticky="w", padx=5, pady=5)
        format_combo = ttk.Combobox(options_frame, textvariable=self.format_var, values=["docx", "pdf", "text"], width=8, state="readonly")
        format_combo.grid(row=0, column=3, padx=5, pady=5)
        
        # Output directory
        ttk.Label(options_frame, text="Output directory:").grid(row=1, column=2, sticky="w", padx=5, pady=5)
        ttk.Entry(options_frame, textvariable=self.output_dir_var, width=10).grid(row=1, column=3, padx=5, pady=5)
        
        # Add browse button for output directory
        browse_output_btn = ttk.Button(options_frame, text="Browse...", command=self.browse_output_dir)
        browse_output_btn.grid(row=1, column=4, padx=5, pady=5)
        
        # Allow duplicates
        self.allow_duplicates_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(
            options_frame, 
            text="Allow duplicates across quizzes",
            variable=self.allow_duplicates_var
        ).grid(row=2, column=0, columnspan=4, sticky="w", padx=5, pady=5)
        
        # Generate button
        generate_btn = ttk.Button(main_frame, text="Generate Quizzes", command=self.generate_quizzes)
        generate_btn.pack(pady=10)
        
        # Log area
        log_frame = ttk.LabelFrame(main_frame, text="Log")
        log_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.log_text = tk.Text(log_frame, height=10, width=70)
        self.log_text.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        scrollbar = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        scrollbar.pack(fill=tk.Y, side=tk.RIGHT)
        self.log_text.config(yscrollcommand=scrollbar.set)
    
    def browse_file(self):
        filename = filedialog.askopenfilename(
            title="Select Question Bank File",
            filetypes=[
                ("Text Files", "*.txt"),
                ("Word Documents", "*.docx"),
                ("Rich Text Format", "*.rtf"),
                ("All Files", "*.*")
            ]
        )
        if filename:
            self.file_var.set(filename)
    
    def browse_output_dir(self):
        directory = filedialog.askdirectory(
            title="Select Output Directory",
            initialdir=self.output_dir_var.get() if self.output_dir_var.get() else "."
        )
        if directory:
            self.output_dir_var.set(directory)
            self.log_message(f"Output directory set to: {directory}")
    
    def log_message(self, message):
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
    
    def load_questions(self):
        filename = self.file_var.get()
        if not filename:
            messagebox.showerror("Error", "Please select a file first.")
            return
        
        try:
            # Show that we're loading
            self.status_var.set("Loading questions...")
            self.root.update_idletasks()
            
            # Load the questions
            self.question_bank = QuestionBank()
            self.question_bank.load_from_file(filename)
            
            # Update status
            question_count = self.question_bank.get_size()
            self.status_var.set(f"Loaded {question_count} questions")
            self.log_message(f"Successfully loaded {question_count} questions from {filename}")
        
        except Exception as e:
            self.status_var.set("Error loading questions")
            messagebox.showerror("Error", f"Failed to load questions: {str(e)}")
            self.log_message(f"Error: {str(e)}")
    
    def generate_quizzes(self):
        if not self.question_bank:
            messagebox.showerror("Error", "Please load a question bank first.")
            return
        
        try:
            # Get parameters
            num_quizzes = self.num_quizzes_var.get()
            questions_per_quiz = self.questions_per_quiz_var.get()
            allow_duplicates = self.allow_duplicates_var.get()
            output_format = self.format_var.get()
            output_dir = self.output_dir_var.get()
            
            # Generate quizzes
            self.log_message(f"Generating {num_quizzes} quizzes with {questions_per_quiz} questions each...")
            self.root.update_idletasks()
            
            randomizer = QuizRandomizer(self.question_bank)
            quizzes, metadata = randomizer.create_quizzes(
                num_quizzes,
                questions_per_quiz,
                allow_duplicates,
                output_format,
                output_dir
            )
            
            # Log results
            result_message = (
                f"\nQuiz Generation Summary:\n"
                f"------------------------\n"
                f"Number of quizzes: {metadata['num_quizzes']}\n"
                f"Questions per quiz: {metadata['questions_per_quiz']}\n"
                f"Total unique questions used: {metadata['total_questions_used']}\n"
            )
            
            if metadata.get("quizzes_with_extra_questions"):
                extra = ", ".join(map(str, metadata["quizzes_with_extra_questions"]))
                result_message += f"Quizzes with extra questions: {extra}\n"
            
            if metadata.get("duplicate_stats"):
                result_message += f"Questions used in multiple quizzes: {len(metadata['duplicate_stats'])}\n"
            
            result_message += f"\nQuizzes saved to the '{output_dir}' directory"
            
            self.log_message(result_message)
            
            if output_format == "pdf":
                self.log_message("\nNote: If PDF generation failed due to encoding issues, text files were created instead.")
                
            messagebox.showinfo("Success", f"Successfully generated {num_quizzes} quizzes!")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate quizzes: {str(e)}")
            self.log_message(f"Error: {str(e)}")


def main():
    root = tk.Tk()
    app = SimpleQuizRandomizerGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()