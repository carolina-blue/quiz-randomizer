import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def extract_and_create_quiz(input_file, output_file, title="Quiz"):
    """Extract questions from a DOCX file and create a new DOCX with proper formatting"""
    # Extract the questions first
    questions = extract_questions_from_docx(input_file)
    
    # Create the output quiz document
    create_formatted_docx(questions, output_file, title)
    
    return len(questions)

def extract_questions_from_docx(filename):
    """Extract questions with proper formatting from a DOCX file"""
    doc = docx.Document(filename)
    
    # Print all paragraphs with their run formatting for debugging
    print("DEBUG: Document content with formatting")
    for i, para in enumerate(doc.paragraphs):
        print(f"Paragraph {i}: '{para.text}'")
        for j, run in enumerate(para.runs):
            print(f"  Run {j}: '{run.text}', Bold: {run.bold}")
    
    questions = []
    current_question = None
    current_options = []
    current_feedback = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Empty paragraph means end of current question
            if current_question:
                questions.append({
                    'text': current_question,
                    'options': current_options.copy(),  # Make sure to copy the list
                    'feedback': current_feedback
                })
                current_question = None
                current_options = []
                current_feedback = None
        elif text.startswith('Answer Feedback:'):
            current_feedback = text
        elif any(text.startswith(f"{chr(97 + i)})") for i in range(26)):
            # This is an option
            current_options.append(text)
        elif not current_question:
            current_question = text
    
    # Don't forget the last question
    if current_question:
        questions.append({
            'text': current_question,
            'options': current_options.copy(),
            'feedback': current_feedback
        })
    
    return questions

def create_formatted_docx(questions, output_file, title="Quiz"):
    """Create a properly formatted DOCX file from the extracted questions"""
    doc = docx.Document()
    
    # Add title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Add some space
    
    # Add questions
    for i, question in enumerate(questions, 1):
        # Question text
        question_paragraph = doc.add_paragraph()
        question_number = question_paragraph.add_run(f"{i}. ")
        question_number.bold = True
        question_paragraph.add_run(question['text'])
        
        # Add options - This part needs to be fixed
        if 'options' in question:
            for option in question['options']:
                option_paragraph = doc.add_paragraph()
                option_paragraph.paragraph_format.left_indent = Pt(20)  # Indent options
                
                # Get the option text
                if isinstance(option, dict):
                    option_text = option.get('text', '')
                else:
                    option_text = option
                    
                # Add the option text
                option_paragraph.add_run(option_text)
        
        # Add space between questions
        doc.add_paragraph()
    
    # Print debug info before saving
    print(f"Saving formatted document with {len(questions)} questions to {output_file}")
    
    # Save the document
    doc.save(output_file)
