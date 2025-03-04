import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re

def create_quiz_from_docx(input_file, output_file, title="Quiz"):
    """
    Create a formatted quiz document from input DOCX with bold answers.
    
    This function handles bold formatting properly by checking each run's bold property.
    """
    # First, extract the questions with their formatting
    questions = extract_questions(input_file)
    
    # Then create the output document
    create_formatted_document(questions, output_file, title)
    
    return len(questions)

def extract_questions(docx_file):
    """Extract questions from a DOCX file, preserving formatting"""
    doc = docx.Document(docx_file)
    
    questions = []
    current_question = None
    current_options = []
    current_feedback = None
    
    in_option_block = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            # Empty paragraph - end current question
            if current_question:
                questions.append({
                    'question': current_question,
                    'options': current_options,
                    'feedback': current_feedback
                })
                current_question = None
                current_options = []
                current_feedback = None
                in_option_block = False
            continue
        
        # Check for option format (a), b), etc. or numbered format (1., 2., etc.)
        # Word sometimes uses special formatting for numbered lists, so we need to check
        # the paragraph's style as well as the content
        is_letter_option = bool(re.match(r'^[a-z]\)', text))
        is_number_option = bool(re.match(r'^\d+\.', text))
        
        # Check Word's list formatting - look for list style or numbering
        has_list_style = False
        if hasattr(para, 'style') and para.style and 'List' in para.style.name:
            has_list_style = True
        elif hasattr(para, 'paragraph_format') and para.paragraph_format.first_line_indent:
            # Word often indents list items
            has_list_style = True
        
        # If this paragraph has numbering applied, it's likely a list item
        has_numbering = hasattr(para, '_element') and para._element.pPr and para._element.pPr.numPr
        
        is_option = is_letter_option or is_number_option or has_list_style or has_numbering
        
        if text.startswith('Answer Feedback:'):
            # Answer feedback line
            current_feedback = text
        elif is_option:
            # This is an option (either letter or number format)
            in_option_block = True
            
            # For number options or Word's list formatting, convert to letter format
            if is_number_option:
                # Standard format with visible number
                number_match = re.match(r'^(\d+)\.\s+(.+)$', text)
                if number_match:
                    number = int(number_match.group(1))
                    content = number_match.group(2)
                    
                    # Convert to letter format (1 -> a, 2 -> b, etc.)
                    letter = chr(96 + number) if 1 <= number <= 26 else chr(96 + (number % 26))
                    text = f"{letter}) {content}"
            elif has_list_style or has_numbering:
                # Word's auto-numbering might not show in the text content
                # Determine number based on position in current options list
                number = len(current_options) + 1  # 1-based index
                content = text
                
                # Convert to letter format
                letter = chr(96 + number) if 1 <= number <= 26 else chr(96 + (number % 26))
                text = f"{letter}) {content}"
            
            # Carefully extract formatting by checking each run
            formatted_parts = []
            for run in para.runs:
                formatted_parts.append({
                    'text': run.text,
                    'bold': run.bold is True  # Explicitly check for True
                })
            
            current_options.append({
                'text': text,
                'formatted_parts': formatted_parts,
                'has_bold': any(part['bold'] for part in formatted_parts)
            })
        elif not current_question and not in_option_block:
            # This must be a question
            current_question = text
    
    # Add the last question if there is one
    if current_question:
        questions.append({
            'question': current_question,
            'options': current_options,
            'feedback': current_feedback
        })
    
    return questions

def create_formatted_document(questions, output_file, title="Quiz"):
    """Create a formatted document with bold answers"""
    doc = docx.Document()
    
    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add a blank line after title
    doc.add_paragraph()
    
    # Add each question
    for i, q in enumerate(questions, 1):
        # Question text with number
        question_para = doc.add_paragraph()
        num_run = question_para.add_run(f"{i}. ")
        num_run.bold = True
        question_para.add_run(q['question'])
        
        # Add each option with proper formatting
        for opt in q['options']:
            option_para = doc.add_paragraph()
            option_para.paragraph_format.left_indent = Pt(20)
            
            # Preserve exact formatting from the original
            if 'formatted_parts' in opt:
                for part in opt['formatted_parts']:
                    run = option_para.add_run(part['text'])
                    if part['bold']:
                        run.bold = True
            else:
                # Fallback if no formatting info
                option_para.add_run(opt['text'])
        
        # Add feedback if present
        if q['feedback']:
            feedback_para = doc.add_paragraph()
            feedback_para.paragraph_format.left_indent = Pt(20)
            feedback_run = feedback_para.add_run(q['feedback'])
            feedback_run.italic = True
        
        # Add space between questions
        doc.add_paragraph()
    
    # Save the document
    doc.save(output_file)
    print(f"Successfully created {output_file} with proper formatting")

# Simple test function
def test():
    """Test the formatting with a sample document"""
    # Create a sample DOCX with bold formatting
    doc = docx.Document()
    
    # Question 1
    doc.add_paragraph('What is the capital of Italy?')
    
    # Options with one bold answer (Rome)
    p = doc.add_paragraph()
    p.add_run('a) Madrid')
    
    p = doc.add_paragraph()
    p.add_run('b) Berlin')
    
    p = doc.add_paragraph()
    p.add_run('c) ')
    p.add_run('Rome').bold = True
    
    p = doc.add_paragraph()
    p.add_run('d) Athens')
    
    # Feedback
    doc.add_paragraph('Answer Feedback: Rome is the capital of Italy')
    doc.add_paragraph()
    
    # Question 2
    doc.add_paragraph('Which planet is largest?')
    
    # Options with one bold answer (Jupiter)
    p = doc.add_paragraph()
    p.add_run('a) Earth')
    
    p = doc.add_paragraph()
    p.add_run('b) Mars')
    
    p = doc.add_paragraph()
    p.add_run('c) ')
    p.add_run('Jupiter').bold = True
    
    p = doc.add_paragraph()
    p.add_run('d) Saturn')
    
    # Feedback
    doc.add_paragraph('Answer Feedback: Jupiter is the largest planet')
    
    # Save the test file
    test_file = 'test_formatting.docx'
    doc.save(test_file)
    print(f"Created test file: {test_file}")
    
    # Create output directory if needed
    os.makedirs("quizzes", exist_ok=True)
    
    # Process the test file
    create_quiz_from_docx(
        input_file=test_file,
        output_file='quizzes/formatted_quiz.docx',
        title="Test Quiz with Formatting"
    )
    
    print("Test complete - check quizzes/formatted_quiz.docx")

if __name__ == "__main__":
    test()