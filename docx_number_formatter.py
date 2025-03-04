import docx
import re
import os

class Question:
    """Simple question class for testing"""
    def __init__(self, text, type="unknown", options=None, feedback=None):
        self.text = text
        self.type = type
        self.options = options or []
        self.feedback = feedback
    
    def __str__(self):
        result = self.text
        for opt in self.options:
            result += f"\n  {opt}"
        if self.feedback:
            result += f"\n  {self.feedback}"
        return result

def create_test_file():
    """Create a test DOCX file with both letter and number formatting"""
    doc = docx.Document()
    
    # Question 1 - Letter format
    doc.add_paragraph('Question 1: What is the capital of Italy?')
    
    # Letter formatted options
    p = doc.add_paragraph()
    p.add_run('a) Madrid')
    
    p = doc.add_paragraph()
    p.add_run('b) Berlin')
    
    p = doc.add_paragraph()
    p.add_run('c) ')
    p.add_run('Rome').bold = True
    
    p = doc.add_paragraph()
    p.add_run('d) Athens')
    
    doc.add_paragraph('Answer Feedback: Rome is the capital of Italy')
    doc.add_paragraph()  # Blank line
    
    # Question 2 - Number format
    doc.add_paragraph('Question 2: Which planet is the largest in our solar system?')
    
    # Number formatted options
    p = doc.add_paragraph()
    p.add_run('1. Earth')
    
    p = doc.add_paragraph()
    p.add_run('2. ')
    p.add_run('Jupiter').bold = True
    
    p = doc.add_paragraph()
    p.add_run('3. Mars')
    
    p = doc.add_paragraph()
    p.add_run('4. Saturn')
    
    doc.add_paragraph('Answer Feedback: Jupiter is the largest planet')
    
    # Save the file
    os.makedirs("quizzes", exist_ok=True)
    test_file = 'quizzes/mixed_format_test.docx'
    doc.save(test_file)
    print(f"Created {test_file} with both letter and number formatting")
    
    return test_file

def load_docx_questions(filename):
    """Load questions from a DOCX file, handling both letter and number options"""
    doc = docx.Document(filename)
    
    # We'll collect questions in a list
    questions = []
    
    # Current question being processed
    current_question = None
    current_options = []
    current_feedback = None
    
    # Process each paragraph
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Skip empty paragraphs
        if not text:
            # End of current question
            if current_question:
                question_type = "free-response"
                if current_options:
                    question_type = "multiple-choice"
                elif current_question.lower().find("true/false") > -1:
                    question_type = "true-false"
                    current_options = ["True", "False"]
                
                questions.append(Question(current_question, question_type, current_options, current_feedback))
                print(f"Added question: {current_question}")
                
                # Reset for next question
                current_question = None
                current_options = []
                current_feedback = None
            continue
        
        # Check for answer feedback
        if text.startswith("Answer Feedback:"):
            current_feedback = text
            continue
        
        # Check for letter-style options (a), b), etc.)
        letter_match = re.match(r'^([a-z]\))\s+(.+)$', text)
        if letter_match:
            letter = letter_match.group(1)
            content = letter_match.group(2)
            
            # Check for bold formatting
            has_bold = False
            for run in para.runs:
                if run.bold is True:
                    has_bold = True
                    bold_text = run.text
                    if bold_text in content:
                        # Mark the bold part
                        content = content.replace(bold_text, f"**{bold_text}**")
                        break
            
            option_text = f"{letter} {content}"
            current_options.append(option_text)
            continue
        
        # Check for number-style options (1., 2., etc.)
        number_match = re.match(r'^(\d+)\.\s+(.+)$', text)
        if number_match:
            number = int(number_match.group(1))
            content = number_match.group(2)
            
            # Convert to letter format (1 -> a, 2 -> b, etc.)
            letter = chr(96 + number) if 1 <= number <= 26 else chr(96 + (number % 26))
            letter_prefix = f"{letter}) "
            
            # Check for bold formatting
            has_bold = False
            for run in para.runs:
                if run.bold is True:
                    has_bold = True
                    bold_text = run.text
                    
                    # If the bold text is in the content
                    if bold_text in content:
                        # Mark the bold part
                        content = content.replace(bold_text, f"**{bold_text}**")
                        break
            
            option_text = f"{letter_prefix}{content}"
            current_options.append(option_text)
            continue
        
        # If we got here, it must be a question text
        if not current_question:
            current_question = text
    
    # Don't forget to add the last question
    if current_question:
        question_type = "free-response"
        if current_options:
            question_type = "multiple-choice"
        elif current_question.lower().find("true/false") > -1:
            question_type = "true-false"
            current_options = ["True", "False"]
        
        questions.append(Question(current_question, question_type, current_options, current_feedback))
        print(f"Added question: {current_question}")
    
    return questions

def main():
    # Create a test file
    test_file = create_test_file()
    
    # Load questions from the test file
    questions = load_docx_questions(test_file)
    
    # Print the questions
    print("\nLoaded Questions:")
    for i, q in enumerate(questions, 1):
        print(f"\nQuestion {i}: {q}")
    
    print("\nThis code successfully handles both letter (a), b)) and number (1., 2.) formatted options.")
    print("Copy the load_docx_questions function into your quiz_randomizer.py file's _load_from_docx method.")

if __name__ == "__main__":
    main()
